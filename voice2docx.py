import tkinter.filedialog
import json
from docx import Document

### Convert results of voice recognition with vosk into json
### 
### @param  voice A list including strings of voice recognition with vosk.
### @return A string of json format.
def voice2json(voice):
    period = ",{'conf': 0.5,'end': 0,'start': 0,'word':'。'},"
    voice = period.join(voice)
    voice = voice.replace(",None", "").replace("\n", "").replace("]", "").replace("[", "").replace("\'", "\"")
    voice_json = '[' + voice + ']'
    return voice_json

### Add word in a word paragraph with a confidential value.
### 
### @param  paragraph  A word paragraph object
### @param  word       A string
### @param  conf       A numeric (0-1)
def add_word_with_conf(paragraph, word, conf):
    run = paragraph.add_run()
    run.text = word
    # if conf < 0.7:
    if conf < 0.33:
        run.font.underline = True
    # if conf < 1:
    if conf < 0.66:
        run.font.bold = True
    return 1

### Write voice dictionary in a docx file.
### 
### @param  voice_dict  A dictionary including result of voice recognition. 
###                     keys: "word" and "conf", values: string and numeric (0-1). 
### @param  path        A string of txt or wav path
### @return A string of docx path
def voice_dict2docx(voice_dict, path):
    docx = path.removesuffix(".txt").removesuffix(".wav") + ".docx"
    document = Document()
    para = document.add_paragraph('')
    for vd in voice_dict:
        add_word_with_conf(para, vd["word"], vd["conf"])
    document.save(docx)
    print("-" * 60 + "\n")
    print("Saved in " + docx + "\n")
    print("-" * 60 + "\n")
    return docx

def voice_dict2txt(voice_dict, path):
    txt = path.removesuffix(".txt").removesuffix(".wav") + "_plain.txt"
    voice_text = ""
    for vd in voice_dict:
        voice_text = voice_text + vd["word"]
    open(txt, mode = "w", encoding = "utf-8").write(voice_text)
    print("-" * 60 + "\n")
    print("Saved in " + txt + "\n")
    print("-" * 60 + "\n")
    return txt

# 
if __name__=="__main__":
    # read txt
    txt = tkinter.filedialog.askopenfilename(filetypes=[("txt", "*.txt")])
    with open(txt, mode = "r", encoding = "utf-8") as f:
       voice = f.readlines()
    # convert string into dictionary
    voice_json = voice2json(voice)
    voice_dict = json.loads(voice_json)
    # write txt
    voice_dict2txt(voice_dict, txt)
    # write docx
    docx = voice_dict2docx(voice_dict, txt)
